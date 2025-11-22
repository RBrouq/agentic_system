from docx import Document
from fpdf import FPDF
import tempfile

def save_docx(title: str, content: str) -> str:
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_paragraph(content)
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp.name)
    return tmp.name

def save_pdf(title: str, content: str) -> str:
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", "B", 16)
    pdf.multi_cell(0, 10, title)
    pdf.ln(5)
    pdf.set_font("Arial", "", 12)
    pdf.multi_cell(0, 8, content)
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
    pdf.output(tmp.name)
    return tmp.name



