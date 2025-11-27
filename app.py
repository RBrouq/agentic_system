from typing import Optional

import textwrap
from tempfile import NamedTemporaryFile

from fastapi.responses import FileResponse
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas

from fastapi import FastAPI, Request, Form
from fastapi.responses import HTMLResponse, JSONResponse
from fastapi.templating import Jinja2Templates
from fastapi.staticfiles import StaticFiles
from dotenv import load_dotenv

from src.runner import run_essay_graph

# Load env vars
load_dotenv()

app = FastAPI(title="Essay Agent – HITL Workflow")

# Static + templates (adapt paths to your project)
app.mount("/static", StaticFiles(directory="static"), name="static")
templates = Jinja2Templates(directory="templates")


def _safe_filename(title: str, ext: str) -> str:
    # Very simple slug
    base = "".join(
        c if c.isalnum() or c in ("-", "_") else "_"
        for c in title.strip().replace(" ", "_")
    )
    if not base:
        base = "essay"
    return f"{base}.{ext}"


@app.get("/", response_class=HTMLResponse)
async def index(request: Request):
    """
    Home page.
    """
    return templates.TemplateResponse(
        "index.html",
        {
            "request": request,
        },
    )


@app.post("/api/run")
async def run_agent(
    prompt: str = Form(...),
    thread_id: Optional[str] = Form(None),

    clarification_answers: Optional[str] = Form(None),
    plan_feedback: Optional[str] = Form(None),
    draft_feedback_human: Optional[str] = Form(None),
    draft_approved: Optional[str] = Form(None),  # will be "true" or None
    final_feedback: Optional[str] = Form(None),

    skip_clarification: Optional[str] = Form(None),  # "on" if checked
    skip_plan_review: Optional[str] = Form(None),
    skip_draft_review: Optional[str] = Form(None),
):
    """
    Run the LangGraph workflow for a given user prompt, with optional HITL inputs.
    """

    # Convert draft_approved checkbox to bool or None
    draft_approved_bool: Optional[bool] = None
    if draft_approved is not None:
        # if checkbox is checked, HTML sends "on" (or "true" from JS)
        draft_approved_bool = True

    # Convert skip checkboxes to bools
    skip_clarification_bool = bool(skip_clarification)
    skip_plan_review_bool = bool(skip_plan_review)
    skip_draft_review_bool = bool(skip_draft_review)

    try:
        result = run_essay_graph(
            prompt,
            thread_id=thread_id,
            clarification_answers=clarification_answers,
            plan_feedback=plan_feedback,
            draft_feedback_human=draft_feedback_human,
            draft_approved=draft_approved_bool,
            final_feedback=final_feedback,
            skip_clarification=skip_clarification_bool,
            skip_plan_review=skip_plan_review_bool,
            skip_draft_review=skip_draft_review_bool,
        )
        return JSONResponse(
            {
                "thread_id": result.get("thread_id"),
                "mode": result.get("mode"),
                "topic": result.get("topic"),
                "instructions": result.get("instructions"),
                "clarification_questions": result.get("clarification_questions"),
                "clarification_answers": result.get("clarification_answers"),
                "plan": result.get("plan"),
                "plan_validated": result.get("plan_validated"),
                "research_notes": result.get("research_notes"),
                "draft": result.get("draft"),
                "critique": result.get("critique"),
                "final_draft": result.get("final_draft"),
                "answer": result.get("answer"),
                "saved": result.get("saved"),
                "final_approved": result.get("final_approved"),
            }
        )
    except Exception as e:
        return JSONResponse(
            {"error": str(e)},
            status_code=500,
        )

@app.post("/api/export/docx")
async def export_docx(
    answer: str = Form(...),
    topic: str = Form("Essay"),
) -> FileResponse:
    """
    Create a .docx file from the essay answer and return it.
    """
    with NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        doc = Document()
        doc.add_heading(topic, level=1)
        doc.add_paragraph("")  # blank line

        for line in answer.splitlines():
            if line.strip():
                doc.add_paragraph(line)
            else:
                doc.add_paragraph("")  # keep blank lines

        doc.save(tmp.name)
        filename = _safe_filename(topic, "docx")

    return FileResponse(
        path=tmp.name,
        media_type=(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
        filename=filename,
    )


@app.post("/api/export/pdf")
async def export_pdf(
    answer: str = Form(...),
    topic: str = Form("Essay"),
) -> FileResponse:
    """
    Create a simple A4 PDF from the essay answer and return it.
    """
    with NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
        c = canvas.Canvas(tmp.name, pagesize=A4)
        width, height = A4

        # Title
        text_obj = c.beginText(50, height - 50)
        text_obj.setFont("Helvetica-Bold", 16)
        text_obj.textLine(topic)
        text_obj.moveCursor(0, -20)

        # Body
        text_obj.setFont("Helvetica", 11)
        max_chars = 90  # quick-and-dirty wrapping per line
        for line in answer.splitlines():
            if not line.strip():
                text_obj.textLine("")  # blank line
                continue

            for chunk in textwrap.wrap(line, max_chars):
                # If we're too low on the page, start a new page
                if text_obj.getY() < 50:
                    c.drawText(text_obj)
                    c.showPage()
                    text_obj = c.beginText(50, height - 50)
                    text_obj.setFont("Helvetica", 11)
                text_obj.textLine(chunk)

        c.drawText(text_obj)
        c.showPage()
        c.save()

        filename = _safe_filename(topic, "pdf")

    return FileResponse(
        path=tmp.name,
        media_type="application/pdf",
        filename=filename,
    )
